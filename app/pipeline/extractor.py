"""Layer 1: Document extraction with waterfall fallback.

Strategy:
  detect file type -> try primary extractor -> validate -> fallback -> OCR
"""
from __future__ import annotations

import io
import subprocess
import tempfile
import time
from dataclasses import dataclass
from pathlib import Path

import chardet
from loguru import logger

from app.services import ocr
from app.utils.text_cleaner import clean_text, is_usable, quality_score


@dataclass
class ExtractedDocument:
    raw_text: str
    file_type: str
    extractor_used: str
    ocr_used: bool
    quality_score: float
    page_count: int
    extraction_time_ms: int


SUPPORTED_TYPES = {"pdf", "docx", "doc", "txt", "rtf"}


def detect_file_type(filename: str, content: bytes) -> str:
    """Detect file type from extension + magic bytes. Returns lowercase ext or 'unknown'."""
    ext = Path(filename).suffix.lower().lstrip(".")

    if ext in SUPPORTED_TYPES:
        # Cross-check against magic bytes for the common cases
        if ext == "pdf" and not content.startswith(b"%PDF"):
            # maybe something else
            pass
        return ext

    # Magic byte fallback
    if content.startswith(b"%PDF"):
        return "pdf"
    if content[:4] == b"PK\x03\x04":
        # DOCX / ZIP
        return "docx"
    if content.startswith(b"\xd0\xcf\x11\xe0"):
        return "doc"
    if content.startswith(b"{\\rtf"):
        return "rtf"

    return ext or "unknown"


def extract_document(filename: str, content: bytes) -> ExtractedDocument:
    """Main entry point: extract text from a document with fallback strategy."""
    t0 = time.time()
    file_type = detect_file_type(filename, content)
    logger.info("extraction_started", filename=filename, file_type=file_type)

    if file_type == "pdf":
        result = _extract_pdf(content)
    elif file_type == "docx":
        result = _extract_docx(content)
    elif file_type == "doc":
        result = _extract_doc(content)
    elif file_type in ("txt", "rtf"):
        result = _extract_text_like(content, file_type)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

    raw_text, extractor_used, ocr_used, page_count = result
    raw_text = clean_text(raw_text)
    score = quality_score(raw_text)
    elapsed = int((time.time() - t0) * 1000)

    logger.info(
        "extraction_complete",
        file_type=file_type,
        extractor=extractor_used,
        ocr_used=ocr_used,
        text_length=len(raw_text),
        quality_score=round(score, 3),
        extraction_time_ms=elapsed,
    )
    return ExtractedDocument(
        raw_text=raw_text,
        file_type=file_type,
        extractor_used=extractor_used,
        ocr_used=ocr_used,
        quality_score=score,
        page_count=page_count,
        extraction_time_ms=elapsed,
    )


# -------------------- PDF --------------------

def _extract_pdf(content: bytes) -> tuple[str, str, bool, int]:
    # Try pdfplumber first (best for columnar layouts)
    text, pages = _pdf_pdfplumber(content)
    if text and is_usable(clean_text(text)):
        return text, "pdfplumber", False, pages

    # Fallback: PyMuPDF
    text2, pages2 = _pdf_pymupdf(content)
    if text2 and is_usable(clean_text(text2)):
        return text2, "pymupdf", False, pages2

    # Both failed or low quality → OCR
    logger.info("pdf_text_low_quality_falling_back_to_ocr")
    ocr_text = ocr.ocr_pdf(content)
    return ocr_text, "ocr", True, pages or pages2 or 0


def _pdf_pdfplumber(content: bytes) -> tuple[str, int]:
    try:
        import pdfplumber
    except ImportError:
        return "", 0
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            page_texts = []
            for page in pdf.pages:
                page_texts.append(_pdfplumber_page_text(page))
            return "\n\n".join(page_texts), len(pdf.pages)
    except Exception as e:
        logger.warning("pdfplumber_failed", error=str(e))
        return "", 0


def _pdfplumber_page_text(page) -> str:
    """Extract text from a pdfplumber page with column-aware ordering.

    Groups words by approximate column, then emits them top-to-bottom.
    """
    words = page.extract_words(x_tolerance=2, y_tolerance=2, keep_blank_chars=False)
    if not words:
        return page.extract_text() or ""

    # Detect columns by clustering x0 values
    page_width = page.width or 612
    mid = page_width / 2
    # simple 2-column detection: split on mid if a reasonable share of words cross it
    left_words = [w for w in words if w["x1"] <= mid + 10]
    right_words = [w for w in words if w["x0"] >= mid - 10]
    crossing = len(words) - len(left_words) - len(right_words)
    is_multi_column = (
        len(left_words) > 10
        and len(right_words) > 10
        and crossing < 0.2 * len(words)
    )

    if is_multi_column:
        left_sorted = sorted(left_words, key=lambda w: (round(w["top"], 1), w["x0"]))
        right_sorted = sorted(right_words, key=lambda w: (round(w["top"], 1), w["x0"]))
        left_text = _words_to_lines(left_sorted)
        right_text = _words_to_lines(right_sorted)
        return left_text + "\n\n" + right_text

    words_sorted = sorted(words, key=lambda w: (round(w["top"], 1), w["x0"]))
    return _words_to_lines(words_sorted)


def _words_to_lines(words: list[dict]) -> str:
    if not words:
        return ""
    lines: list[list[str]] = []
    current_top = None
    current_line: list[str] = []
    for w in words:
        top = round(w["top"], 0)
        if current_top is None or abs(top - current_top) <= 3:
            current_line.append(w["text"])
            current_top = current_top if current_top is not None else top
        else:
            lines.append(current_line)
            current_line = [w["text"]]
            current_top = top
    if current_line:
        lines.append(current_line)
    return "\n".join(" ".join(line) for line in lines)


def _pdf_pymupdf(content: bytes) -> tuple[str, int]:
    try:
        import fitz
    except ImportError:
        return "", 0
    try:
        doc = fitz.open(stream=content, filetype="pdf")
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text("text"))
        n = len(doc)
        doc.close()
        return "\n\n".join(text_parts), n
    except Exception as e:
        logger.warning("pymupdf_failed", error=str(e))
        return "", 0


# -------------------- DOCX --------------------

def _extract_docx(content: bytes) -> tuple[str, str, bool, int]:
    try:
        from docx import Document
    except ImportError:
        logger.error("python_docx_not_installed")
        return "", "docx_failed", False, 0

    try:
        doc = Document(io.BytesIO(content))
    except Exception as e:
        logger.warning("docx_open_failed", error=str(e))
        return "", "docx_failed", False, 0

    parts: list[str] = []

    # Headers
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            if para.text.strip():
                parts.append(para.text)

    # Body: traverse document.body in order so paragraphs/tables are in reading order
    body = doc.element.body
    from docx.oxml.ns import qn

    for child in body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            # paragraph
            for para in doc.paragraphs:
                if para._element is child:
                    if para.text.strip():
                        parts.append(para.text)
                    break
        elif tag == qn("w:tbl"):
            for table in doc.tables:
                if table._element is child:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            parts.append(row_text)
                    break

    text = "\n".join(parts)
    return text, "python-docx", False, 1


# -------------------- DOC --------------------

def _extract_doc(content: bytes) -> tuple[str, str, bool, int]:
    """Convert .doc -> .docx via LibreOffice headless, then extract."""
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            src = Path(tmpdir) / "input.doc"
            src.write_bytes(content)
            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    tmpdir,
                    str(src),
                ],
                capture_output=True,
                timeout=60,
                check=False,
            )
            if result.returncode != 0:
                logger.warning("libreoffice_convert_failed", stderr=result.stderr.decode(errors="ignore"))
                return "", "doc_failed", False, 0
            docx_path = Path(tmpdir) / "input.docx"
            if not docx_path.exists():
                return "", "doc_failed", False, 0
            docx_content = docx_path.read_bytes()
        text, _, _, pages = _extract_docx(docx_content)
        return text, "libreoffice+python-docx", False, pages
    except FileNotFoundError:
        logger.error("libreoffice_not_installed")
        return "", "doc_failed", False, 0
    except subprocess.TimeoutExpired:
        logger.error("libreoffice_timeout")
        return "", "doc_failed", False, 0


# -------------------- TXT / RTF --------------------

def _extract_text_like(content: bytes, file_type: str) -> tuple[str, str, bool, int]:
    encoding = chardet.detect(content).get("encoding") or "utf-8"
    try:
        text = content.decode(encoding, errors="replace")
    except (LookupError, TypeError):
        text = content.decode("utf-8", errors="replace")

    if file_type == "rtf":
        text = _strip_rtf(text)

    return text, f"text-{file_type}", False, 1


def _strip_rtf(text: str) -> str:
    """Very basic RTF control-word stripper."""
    import re

    text = re.sub(r"\\[a-z]+-?\d*\s?", "", text)
    text = re.sub(r"[{}]", "", text)
    text = re.sub(r"\\'[0-9a-fA-F]{2}", "", text)
    return text


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("usage: python -m app.pipeline.extractor <file>")
        sys.exit(1)
    p = Path(sys.argv[1])
    result = extract_document(p.name, p.read_bytes())
    print(f"extractor: {result.extractor_used}")
    print(f"ocr_used:  {result.ocr_used}")
    print(f"quality:   {result.quality_score:.3f}")
    print(f"length:    {len(result.raw_text)}")
    print("----")
    print(result.raw_text[:2000])
